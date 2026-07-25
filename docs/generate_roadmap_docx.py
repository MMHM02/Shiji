"""
食记 V1.0 开发总纲 — Word 文档生成器
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── 页面设置 ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 辅助函数 ──────────────────────────────────────────────

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_colored_heading(text, level, color="2E7D32"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_phase_header(phase_num, phase_name, weeks, color="2E7D32"):
    """添加阶段标题"""
    h = doc.add_heading(f"Phase {phase_num}  {phase_name}", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)

def style_table(table, header_color="2E7D32"):
    """美化表格"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头样式
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9.5)
    # 数据行样式
    for i, row in enumerate(table.rows[1:], 1):
        bg = "F5F5F5" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            set_cell_shading(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

def add_checklist(items, columns=2):
    """添加检查清单表格"""
    table = doc.add_table(rows=len(items) + 1, cols=columns)
    table.style = 'Table Grid'
    header = table.rows[0].cells
    header[0].text = "✅ 检查项"
    header[1].text = "状态"
    for i, item in enumerate(items, 1):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = "☐"
    style_table(table)
    return table


# ══════════════════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════════════════

doc.add_paragraph()  # 空行
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("食  记")
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string("2E7D32")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("AI 饮食记录 · 健康数据看板")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor.from_string("666666")

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f"V1.0 开发总纲  ·  {datetime.date.today().strftime('%Y年%m月%d日')}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor.from_string("999999")

doc.add_paragraph()
doc.add_paragraph()

# 概览表
overview_title = doc.add_paragraph()
overview_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = overview_title.add_run("项目速览")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor.from_string("2E7D32")

overview = doc.add_table(rows=7, cols=2)
overview.style = 'Table Grid'
overview_data = [
    ("项目名称", "食记"),
    ("平台", "Android 8.0+ (API 26+)"),
    ("技术栈", "Kotlin + Jetpack Compose + Room + Hilt"),
    ("AI 厂商", "DeepSeek / Kimi / 通义千问 / GLM + 自定义端点"),
    ("核心理念", "用户自持 API Key · 数据纯本地 · 100% 开源"),
    ("总工期", "18 周 / 7 个 Phase"),
    ("发布方式", "GitHub Release + Google Play"),
]
for i, (k, v) in enumerate(overview_data):
    overview.rows[i].cells[0].text = k
    overview.rows[i].cells[1].text = v
    for cell in overview.rows[i].cells[0:1]:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
style_table(overview)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  路线图总览
# ══════════════════════════════════════════════════════════

add_colored_heading("一、V1.0 阶段总览", 1)

p = doc.add_paragraph()
p.add_run("整个 V1.0 分为 7 个开发阶段，共计 18 周。每个阶段有明确的里程碑，阶段之间按顺序推进，部分任务可并行。").font.size = Pt(10.5)

doc.add_paragraph()

# 总览表
summary = doc.add_table(rows=8, cols=4)
summary.style = 'Table Grid'
headers = ["阶段", "名称", "周期", "里程碑"]
for i, h in enumerate(headers):
    summary.rows[0].cells[i].text = h

phase_overview = [
    ("Phase 0", "项目初始化", "Week 1-2", "空应用跑通 · CI 绿灯"),
    ("Phase 1", "基础架构", "Week 3-4", "3 Tab 导航 · 数据库就绪"),
    ("Phase 2", "饮食记录核心", "Week 5-8", "手动+拍照+语音三通道完整"),
    ("Phase 3", "AI 集成", "Week 9-10", "AI 分析闭环 · 国内大模型接入"),
    ("Phase 4", "健康追踪", "Week 11-12", "体重+运动+Health Connect+首次引导"),
    ("Phase 5", "数据看板", "Week 13-14", "首页+数据Tab+Widget+.fitness导出"),
    ("Phase 6", "打磨与测试", "Week 15-16", "测试达标 · 性能合格"),
]
for i, (phase, name, weeks, milestone) in enumerate(phase_overview, 1):
    row = summary.rows[i]
    row.cells[0].text = phase
    row.cells[1].text = name
    row.cells[2].text = weeks
    row.cells[3].text = milestone
    for cell in [row.cells[0]]:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
style_table(summary)

# 进度条示意
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("进度示意：").font.bold = True
p.add_run("\n■ 核心不可裁剪    ■ 重要尽量保留    ■ 可延后或砍掉").font.size = Pt(9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  各 Phase 详细
# ══════════════════════════════════════════════════════════

add_colored_heading("二、各阶段详细计划", 1)

phases = [
    {
        "num": "0",
        "name": "项目初始化",
        "weeks": "Week 1-2",
        "goal": "空应用可运行，架构骨架就绪，CI/CD 绿灯。",
        "tasks": [
            ("创建 Android 多模块项目", "app/ + core/ + feature/ Gradle 模块", "■"),
            ("配置 Gradle Catalog", "libs.versions.toml 统一依赖管理", "■"),
            ("搭建 Clean Architecture 三层目录", "Presentation → Domain → Data", "■"),
            ("配置 Hilt + Room + Navigation", "DI 容器 + 数据库框架 + 导航框架", "■"),
            ("配置 Material 3 主题", "明/暗色彩方案 + 字体排版 + 形状系统", "■"),
            ("配置 CI/CD（GitHub Actions）", "PR 触发 lint + test + build", "■"),
            ("配置代码规范", "ktlint + detekt 规则", "■"),
            ("国际化框架", "简体中文 strings.xml", "■"),
        ],
        "checklist": [
            "应用可安装到模拟器并显示空白首页",
            "./gradlew lint test build 全部通过",
            "CI 自动触发且绿灯",
        ],
    },
    {
        "num": "1",
        "name": "基础架构",
        "weeks": "Week 3-4",
        "goal": "3 个 Tab 页面可自由导航，数据库全部建好，权限框架和通用组件就绪。",
        "tasks": [
            ("Room 数据库实体 + DAO 定义", "FoodRecord / HealthMetric / UserGoal / AiProvider / CachedFoodItem", "■"),
            ("DataStore 偏好存储封装", "设置项读写（目标值、主题偏好等）", "■"),
            ("3 Tab 导航搭建", "首页 / 数据 / 我的 — BottomNavigationBar", "■"),
            ("通用 UI 组件库", "按钮、卡片、输入框、加载态、骨架屏、错误提示", "■"),
            ("权限管理框架", "相机 + 录音 + Health Connect 权限申请", "■"),
            ("API Key 安全存储", "EncryptedSharedPreferences 封装", "■"),
            ("统一错误处理框架", "Result<T> 封装 + 错误映射到 UI 状态", "■"),
            ("相机模块基础封装", "CameraX 初始化 + 预览 + 拍照回调", "■"),
            ("语音识别接口抽象", "SpeechRecognitionService 接口 + Android 原生实现", "■"),
        ],
        "checklist": [
            "3 个 Tab 页面可自由切换",
            "数据库迁移脚本可顺利执行",
            "权限请求弹窗正常，拒绝后降级可用",
        ],
    },
    {
        "num": "2",
        "name": "饮食记录核心",
        "weeks": "Week 5-8（4 周，最重要阶段）",
        "goal": "手动、拍照、语音三种饮食记录方式全部走通。个人食物库可用。照片 3 天自动清理。",
        "tasks": [
            ("手动添加食物页面", "名称 + 份量 + 单位 + 热量 + 宏量输入", "■"),
            ("个人食物库（从零建立）", "条目=菜名+克重+热量，增删改查+搜索", "■"),
            ("餐次系统", "默认 4 模板（早/午/晚/加餐），可自定义，按时间自动推断", "■"),
            ("饮食日志列表页", "日期切换 + 餐次分组 + 热量小计 + 编辑/删除", "■"),
            ("拍照界面", "CameraX 全屏预览 + 拍照按钮 + 相册入口", "■"),
            ("图片预处理", "压缩≤2MB + 最长边≤2048px + EXIF旋转修正", "■"),
            ("拍照→预览→确认流程", "预览→重拍/使用→确认页", "■"),
            ("照片 3 天自动清理", "WorkManager 定期任务", "■"),
            ("录音界面", "按住说话 + 松开发送 + 实时波形 + 上滑取消", "■"),
            ("ASR 语音识别", "Android SpeechRecognizer 封装 + 实时上屏", "■"),
            ("语音→确认流程", "识别文字→确认页", "■"),
            ("份量单位转换工具", "g/ml/份/碗/个 之间参考换算", "■"),
        ],
        "checklist": [
            "手动添加并保存一条完整饮食记录",
            "食物库可搜索并复用已有条目",
            "拍照→预览→确认→保存 全流程走通",
            "语音按住录音→松手识别→确认保存 全流程走通",
            "饮食日志按日期和餐次正确分组展示",
            "3 天前照片被自动删除",
        ],
    },
    {
        "num": "3",
        "name": "AI 集成",
        "weeks": "Week 9-10",
        "goal": "拍照识食 + 语音解析 AI 分析闭环。DeepSeek / Kimi / 通义千问 / GLM + 自定义端点全部接入。AI 4 种错误场景全覆盖。",
        "tasks": [
            ("AiService 统一接口定义", "analyzeFoodImage / parseFoodDescription / chat / chatStream", "■"),
            ("国内大模型适配器", "DeepSeek / Kimi / 通义千问 / GLM / OpenAI 兼容自定义", "■"),
            ("AiServiceManager（路由）", "按视觉槽位和文本槽位选择适配器", "■"),
            ("AI 响应解析器", "JSON 解析 + 容错 + 格式修正 + 结构化", "■"),
            ("API Key 管理页面", "厂商选择 → Key 输入 → 连接测试 → 加密存储", "■"),
            ("模型槽位配置", "视觉模型 vs 文本模型，分开选择", "■"),
            ("视觉模型不可用处理", "无视觉能力的模型 → 提示 + 引导语音描述路径", "■"),
            ("拍照识食 AI 连通", "拍照→压缩→视觉API→解析→确认页(可逐项编辑/删除)→保存", "■"),
            ("语音解析 AI 连通", "录音→ASR→文本API→解析→确认页→保存", "■"),
            ("AI 错误 4 场景覆盖", "超时提醒 / 解析失败降级 / 单项编辑删除 / 非食物提示", "■"),
            ("用量统计", "计数 + Token 估算 + 费用展示", "■"),
            ("AI 功能首次引导", "提示配置 API Key 的图文引导", "■"),
        ],
        "checklist": [
            "拍照→AI→确认→保存 全流程走通",
            "语音→ASR→AI→确认→保存 全流程走通",
            "API Key 可配置、可测试、加密存储",
            "4 种错误场景均有对应处理",
            "无视觉能力模型时正确提示并引导语音路径",
        ],
    },
    {
        "num": "4",
        "name": "健康追踪",
        "weeks": "Week 11-12",
        "goal": "体重记录+趋势图可用。运动记录+Health Connect 手环数据同步。首次启动 3 步引导可用。",
        "tasks": [
            ("体重记录页面", "卡片形态：大字最新体重 + 缩略折线图，点击进详情页", "■"),
            ("体重趋势图（Vico）", "7天 / 30天 / 90天 可切换", "■"),
            ("运动记录（手动）", "运动类型 + 时长 + 强度 + 估算消耗", "■"),
            ("Health Connect 集成", "读取权限 + 体重 / 卡路里消耗 / 步数 / 心率 / 睡眠", "■"),
            ("HC 不可用降级", "非小米/华为设备静默降级，手动录入不受影响", "■"),
            ("水分摄入记录", "快捷加水 + 每日统计", "■"),
            ("目标设定页面", "身高/体重/目标→自动计算每日热量和宏量", "■"),
            ("首次启动 3 步引导", "Step1 设目标→Step2 配AI→Step3 完成，均可跳过", "■"),
        ],
        "checklist": [
            "体重手动录入，趋势图正确渲染",
            "Health Connect 权限申请+数据读取正常（小米设备）",
            "HC 不可用时手动录入功能毫无影响",
            "首次启动引导可完整走完，跳过也有兜底卡片",
        ],
    },
    {
        "num": "5",
        "name": "数据看板",
        "weeks": "Week 13-14",
        "goal": "首页仪表盘+数据 Tab 完整。桌面 Widget 可用。.fitness 文件导出/导入正常。",
        "tasks": [
            ("首页今日热量圆环", "动画圆环：已摄入/目标 + 百分比", "■"),
            ("首页营养素达标条", "蛋白质/碳水/脂肪 进度条 + 达标/差多少", "■"),
            ("首页快捷入口", "拍照📸/语音🎤/手动✏️ 三按钮卡片", "■"),
            ("首页今日饮食列表", "餐次分组 + 热量小计 + 查看全部", "■"),
            ("数据 Tab 本周概览", "日均热量/蛋白质/碳水 + ✅/⚠️达标状态", "■"),
            ("数据 Tab 热量趋势", "7天折线图", "■"),
            ("数据 Tab 体重趋势", "7天/30天折线图 + 详情入口", "■"),
            ("数据 Tab AI建议卡片流", "基于饮食自动生成3条建议 + 可追问", "■"),
            ("数据 Tab 自由提问入口", "底部输入框 → AI 对话", "■"),
            ("桌面热量圆环 Widget", "2×2 小组件", "■"),
            (".fitness 导出/导入", "标准 JSON（无密码）→ .fitness 后缀", "■"),
        ],
        "checklist": [
            "首页热量圆环正确展示今日进度",
            "数据 Tab 本周概览数值正确",
            "趋势图可切换时间范围",
            "AI 建议卡片自动生成有意义的建议",
            ".fitness 导出→删除→导入→数据恢复 完整流程正常",
            "Widget 可添加到桌面，更新正常",
        ],
    },
    {
        "num": "6",
        "name": "打磨与测试",
        "weeks": "Week 15-16",
        "goal": "测试覆盖率达标（单元≥80%，DAO≥90%），性能合格，暗色/无障碍覆盖，所有错误状态兜底。",
        "tasks": [
            ("单元测试（UseCase/Repository）", "覆盖率 ≥ 80%", "■"),
            ("DAO 测试（Room In-Memory）", "覆盖率 ≥ 90%", "■"),
            ("Compose UI 测试", "核心页面：首页/饮食记录/API配置/体重", "■"),
            ("错误状态全覆盖", "空态/加载中/网络错误/AI异常/HC不可用", "■"),
            ("离线可用性", "断网时手动记录正常，AI功能提示联网", "■"),
            ("照片清理验证", "确认3天前照片被删除", "■"),
            ("暗色模式适配", "所有页面暗色模式正常", "■"),
            ("4.7\"小屏适配", "不溢出，可滚动", "■"),
            ("无障碍适配", "TalkBack可用，对比度WCAG AA", "■"),
            ("性能优化", "冷启动≤1.5s，列表60fps，内存≤200MB", "■"),
        ],
        "checklist": [
            "所有单元测试通过",
            "关键流程 UI 测试通过",
            "离线可用，AI 不可用时降级正常",
            "暗色模式全页面覆盖",
            "性能指标达标",
        ],
    },
]

# 输出每个 Phase
for i, phase in enumerate(phases):
    if i > 0:
        doc.add_page_break()

    # Phase 标题
    color = "2E7D32" if i < 7 else "FF6F00"
    add_phase_header(phase["num"], phase["name"], phase["weeks"], color)

    # 目标
    p = doc.add_paragraph()
    run = p.add_run("🎯 阶段目标：")
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(phase["goal"]).font.size = Pt(11)

    # 工期
    p = doc.add_paragraph()
    run = p.add_run("⏱ 工期：")
    run.font.bold = True
    p.add_run(phase["weeks"])

    doc.add_paragraph()

    # 任务表
    table = doc.add_table(rows=len(phase["tasks"]) + 1, cols=3)
    table.style = 'Table Grid'
    headers = ["#", "任务", "描述"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for j, (task_name, task_desc, priority) in enumerate(phase["tasks"], 1):
        row = table.rows[j]
        row.cells[0].text = str(j)
        row.cells[1].text = task_name
        row.cells[2].text = task_desc
        # 优先级着色
        if priority == "■":
            for run in row.cells[0].paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string("C62828")

    style_table(table)

    doc.add_paragraph()

    # 检查清单
    p = doc.add_paragraph()
    run = p.add_run("✅ 检查点：")
    run.font.bold = True
    run.font.size = Pt(10.5)

    checklist = doc.add_table(rows=len(phase["checklist"]) + 1, cols=2)
    checklist.style = 'Table Grid'
    checklist.rows[0].cells[0].text = "检查项"
    checklist.rows[0].cells[1].text = "状态"
    for j, item in enumerate(phase["checklist"], 1):
        checklist.rows[j].cells[0].text = item
        checklist.rows[j].cells[1].text = "⬜"
    style_table(checklist, "FF6F00")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  Phase 7: 发布准备
# ══════════════════════════════════════════════════════════

add_colored_heading("三、Phase 7 — 发布准备", 1)

p = doc.add_paragraph()
run = p.add_run("🎯 阶段目标：")
run.font.bold = True
p.add_run("食记 V1.0 正式发布到 GitHub Release（及 Google Play）。")

p = doc.add_paragraph()
run = p.add_run("⏱ 工期：")
run.font.bold = True
p.add_run("Week 17-18")

doc.add_paragraph()

release_table = doc.add_table(rows=10, cols=2)
release_table.style = 'Table Grid'
release_tasks = [
    ("应用图标设计", "食记品牌图标，适配各分辨率（mdpi ~ xxxhdpi）"),
    ("首次使用引导润色", "3 步引导 UI 最终打磨，确保无死胡同"),
    ("隐私政策页面", "应用内展示：数据全本地 / 不上传 / API Key 加密存储"),
    ("README 完善", "功能说明 + 截图 + 快速开始 + API Key 获取教程（国内厂商）"),
    ("GitHub Release 准备", "Release Notes + APK 附件"),
    ("Google Play 商店列表", "截图 + 描述 + 功能图（如需上架）"),
    ("应用签名 + ProGuard", "Upload Key + R8 混淆规则 + 测试混淆后 APK"),
    ("内测分发", "GitHub Release APK + Google Play Internal Testing"),
    ("Bug 修复", "内测反馈集中修复，优先级 P0 > P1 > P2"),
]
for j, (k, v) in enumerate(release_tasks, 1):
    release_table.rows[j].cells[0].text = k
    release_table.rows[j].cells[1].text = v
# 表头
release_table.rows[0].cells[0].text = "任务"
release_table.rows[0].cells[1].text = "描述"
style_table(release_table)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  最小可用集
# ══════════════════════════════════════════════════════════

add_colored_heading("四、V1.0 最小可用集（12 周降级方案）", 1, "C62828")

p = doc.add_paragraph()
p.add_run("如果时间严重压缩，以下为绝对不能砍的最小集合：").font.size = Pt(10.5)

doc.add_paragraph()

min_table = doc.add_table(rows=10, cols=2)
min_table.style = 'Table Grid'

min_items = [
    ("✅ 必须保留（P0）", "❌ 可延至 V1.1"),
    ("手动添加食物记录", "多厂商深度适配（保留 DeepSeek + 自定义）"),
    ("个人食物库（从零建立）", "Health Connect 集成"),
    ("拍照识食 + AI 分析", "AI 建议卡片流"),
    ("语音口述 + AI 解析", "桌面 Widget"),
    ("API Key 管理（国内大模型）", "水分/围度记录"),
    ("饮食日志（日期+餐次）", "暗色模式全覆盖"),
    ("首页仪表盘（热量圆环+快捷入口）", "无障碍全面适配"),
    ("体重记录 + 趋势图", "多语言"),
    (".fitness 导出/导入", ""),
]

for j, (keep, drop) in enumerate(min_items):
    min_table.rows[j].cells[0].text = keep
    min_table.rows[j].cells[1].text = drop
    if j == 0:
        set_cell_shading(min_table.rows[j].cells[0], "2E7D32")
        set_cell_shading(min_table.rows[j].cells[1], "C62828")
        for cell in min_table.rows[j].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

style_table(min_table)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  发布后迭代
# ══════════════════════════════════════════════════════════

add_colored_heading("五、发布后迭代规划", 1, "FF6F00")

# V1.1
p = doc.add_paragraph()
run = p.add_run("V1.1 — 体验增强（发布后 4-6 周）")
run.font.bold = True
run.font.size = Pt(12)

v11 = doc.add_table(rows=7, cols=2)
v11.style = 'Table Grid'
v11_items = [
    ("应用内反馈入口", "摇一摇反馈 / 设置页反馈"),
    ("AI 识别准确率优化", "Prompt 迭代 + 多模型对比"),
    ("性能优化", "冷启动加速 + DB 索引优化"),
    ("更多 AI 厂商原生适配", "Kimi / 通义千问 / GLM 原生 API"),
    ("AI 语音识别替代", "可选切换 AI API 做语音转文字"),
    ("自定义营养目标", "生酮/低碳/素食等饮食模式预设"),
]
v11.rows[0].cells[0].text = "功能"
v11.rows[0].cells[1].text = "说明"
for j, (k, v) in enumerate(v11_items, 1):
    v11.rows[j].cells[0].text = k
    v11.rows[j].cells[1].text = v
style_table(v11)

doc.add_paragraph()

# V1.2
p = doc.add_paragraph()
run = p.add_run("V1.2 — 功能增强（发布后 2-3 个月）")
run.font.bold = True
run.font.size = Pt(12)

v12 = doc.add_table(rows=8, cols=2)
v12.style = 'Table Grid'
v12_items = [
    ("食谱推荐引擎", "AI 根据目标/偏好/季节推荐一周食谱"),
    ("饮食模板", "常吃组合餐一键记录"),
    ("自定义菜式嵌套", "菜式 = 原材料组合（番茄炒蛋 = 鸡蛋×2 + 番茄×1）"),
    ("营养元素深分析", "AI 分析微量元素缺口"),
    ("饮食分享卡", "生成精美每日饮食总结图片"),
    ("多设备同步", ".fitness 文件手动同步"),
    ("更多手环数据", "睡眠/心率等 Health Connect 数据展示"),
]
v12.rows[0].cells[0].text = "功能"
v12.rows[0].cells[1].text = "说明"
for j, (k, v) in enumerate(v12_items, 1):
    v12.rows[j].cells[0].text = k
    v12.rows[j].cells[1].text = v
style_table(v12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  风险矩阵
# ══════════════════════════════════════════════════════════

add_colored_heading("六、风险矩阵", 1, "C62828")

risk = doc.add_table(rows=7, cols=4)
risk.style = 'Table Grid'
risk_headers = ["风险", "概率", "影响", "缓解措施"]
for j, h in enumerate(risk_headers):
    risk.rows[0].cells[j].text = h

risk_data = [
    ("AI API 响应不稳定", "中", "高", "多轮重试 + 降级手动输入"),
    ("国内大模型视觉能力差", "中", "高", "视觉槽位默认指向 OpenAI 兼容中转"),
    ("Health Connect 厂商覆盖有限", "中", "中", "手动录入是第一公民，HC 静默降级"),
    ("原生 ASR 中文效果差", "高", "中", "接口已抽象，可切换 AI API 语音转文字"),
    ("用户获取困难", "高", "中", "GitHub 开源引流 + 技术社区推广"),
    ("开发进度延期", "中", "中", "有 12 周最小集降级方案，砍非核心保闭环"),
]
for j, (r, prob, impact, fix) in enumerate(risk_data, 1):
    risk.rows[j].cells[0].text = r
    risk.rows[j].cells[1].text = prob
    risk.rows[j].cells[2].text = impact
    risk.rows[j].cells[3].text = fix
    # 概率/影响着色
    if prob == "高":
        set_cell_shading(risk.rows[j].cells[1], "FFCDD2")
    if impact == "高":
        set_cell_shading(risk.rows[j].cells[2], "FFCDD2")

style_table(risk)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  架构决策一览
# ══════════════════════════════════════════════════════════

add_colored_heading("七、关键架构决策（ADR）一览", 1, "1565C0")

adr = doc.add_table(rows=15, cols=3)
adr.style = 'Table Grid'
adr_headers = ["决策", "结论", "理由"]
for j, h in enumerate(adr_headers):
    adr.rows[0].cells[j].text = h

adr_data = [
    ("V1 目标用户", "GitHub 技术人群 + 国内为主", "GitHub 首发天然筛选；国内大模型 Key 门槛低"),
    ("智能手环同步", "Health Connect（不直连蓝牙）", "一次对接覆盖多品牌；小米/华为已支持"),
    ("数据备份", "本地存储 + JSON 一键导出/导入", "轻量实现；用户自行管理备份位置"),
    ("AI 厂商", "国内大模型 + OpenAI 兼容端点", "DeepSeek/Kimi/Qwen/GLM + 自定义扩展"),
    ("视觉 vs 文本", "分开配置两个槽位", "国内多数模型不支持视觉；视觉走兼容中转"),
    ("语音方案", "原生 ASR + 接口抽象", "V1 快速可用；保留切换到 AI API 的后路"),
    ("食物库", "个人从零建立：菜名+克重+热量", "简单才能坚持；随用户饮食自然增长"),
    ("产品本质", "AI 饮食记录 + 基础健康看板", "不碰训练计划；聚焦核心价值"),
    ("底部导航", "3 Tab：首页 / 数据 / 我的", "AI 顾问融合进数据 Tab 建议卡片流"),
    ("AI 顾问形态", "主动建议卡片 + 自由提问", "解决空白对话页冷启动问题"),
    ("AI 错误处理", "4 场景覆盖", "超时/解析失败/单项错误/非食物"),
    ("首次启动", "3 步引导（可跳过）", "设目标→配AI→完成；跳过有兜底卡片"),
    ("Widget", "V1 P0", "桌面热量圆环；靠视觉可见性维持留存"),
    ("App 名称", "食记", "功能直白、好记、国内友好"),
]
for j, (d, c, r) in enumerate(adr_data, 1):
    adr.rows[j].cells[0].text = d
    adr.rows[j].cells[1].text = c
    adr.rows[j].cells[2].text = r
    for run in adr.rows[j].cells[0].paragraphs[0].runs:
        run.font.bold = True

style_table(adr, "1565C0")

# ══════════════════════════════════════════════════════════
#  保存
# ══════════════════════════════════════════════════════════

output_path = r"e:\Project All\fitness\docs\ShiJi_V1_Development_Plan.docx"
doc.save(output_path)
print(f"Done: {output_path}")
